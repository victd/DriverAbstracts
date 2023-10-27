

import docx

from docx.shared import Pt

doc = docx.Document()

#style = doc.styles['Normal']
#font = style.font
#font.name = 'Courier New'
#font.size = Pt(10)


# CHANGE THIS TEXT FILE FOR THE DIVISION ----------------------------
# are you able to change the format with python and save as pdf as well (using docx2pdf)
# as of Mar 2023, use from fpdf import FPDF, this will write to a pdf file directly and use pdf.add_page() for the page breaks
# export directly to write to a PDF file with page breaks (import FPDF), bypass the word doc, there are also
# its an IT initiative, although safety only asks for it quarterly, active directory
# install python on user's machine or on the server
# blogTO yonge st clair - Employee abstracts as well as driver abstracts - Add page breaks
# Barrie km odo correction, cold condensation
# to facilitate the UI interface, form fields need to be added as in the pdf form, logos and images, tables, lines
# convert special characterss and Greek alphabet to plain text - escape characters excluded, prevent cross site scripts

f = open("Barrie employee abstracts from MTO - Jan 2020.txt", "r")

Lines = f.readlines()

#currline = f.readline()
#doc.add_paragraph(currline)

#count = 0

for line in Lines:
	doc.add_paragraph(line.strip())
	if "END OF RECORD" in line:
		doc.add_page_break()



f.close()


# CHANGE THIS DOCX FILE FOR THE DIVISION ----------------------------

doc.save('Barrie employee abstracts from MTO - Jan 2020test.docx')




